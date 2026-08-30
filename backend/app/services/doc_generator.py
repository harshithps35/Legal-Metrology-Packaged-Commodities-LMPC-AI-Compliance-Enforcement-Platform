"""
LMPC Compliance System — Document / Excel Export Generator

Generates Excel spreadsheets (.xlsx) and Word documents (.docx) for regulatory
notices, legal notices, and bulk compliance auditing reports.
"""

import io
from datetime import datetime, timezone
from typing import Optional


def generate_docx_report(scan_data: dict, output_path: Optional[str] = None) -> bytes:
    """Generate a Microsoft Word (.docx) Legal Metrology Inspection Notice.

    Args:
        scan_data: Dict containing scan details, extracted_fields, violations, and report.
        output_path: Optional file path to save the generated docx.

    Returns:
        bytes of the generated Word document.
    """
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return _generate_csv_fallback(scan_data).encode('utf-8')

    doc = docx.Document()

    # Document Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("DIRECTORATE OF LEGAL METROLOGY\n")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    run_sub = p_title.add_run("OFFICIAL COMPLIANCE AUDIT NOTICE & REPORT\n")
    run_sub.bold = True
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_leg = p_sub.add_run("Issued under the Legal Metrology Act, 2009 & Packaged Commodities Rules, 2011")
    run_leg.italic = True
    run_leg.font.size = Pt(9)
    run_leg.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Metadata Table
    status_str = (scan_data.get('status') or 'REQUIRES_REVIEW').upper()
    score = scan_data.get('compliance_score', 0) or 0

    meta_table = doc.add_table(rows=4, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = True

    meta_data = [
        ("Inspection ID:", f"#{scan_data.get('id', 'N/A')}", "Status / Verdict:", status_str),
        ("Product Name:", scan_data.get('product_name') or 'N/A', "Compliance Score:", f"{score:.1f}%"),
        ("Brand / Packer:", scan_data.get('brand') or 'Unspecified', "Category:", (scan_data.get('category') or 'General').capitalize()),
        ("Audit Date:", datetime.now(timezone.utc).strftime("%d %b %Y, %H:%M UTC"), "Auditing Officer:", (scan_data.get('user') or {}).get('full_name') or 'Inspector'),
    ]

    for row_idx, row_values in enumerate(meta_data):
        for col_idx, val in enumerate(row_values):
            cell = meta_table.cell(row_idx, col_idx)
            cell.text = val
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx % 2 == 0:
                p.runs[0].bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 1. Mandatory Declarations
    h1 = doc.add_heading("1. Statutory Mandatory Declarations Checklist (Rule 6)", level=2)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(6)

    fields_list = scan_data.get('extracted_fields', [])
    field_table = doc.add_table(rows=1, cols=4)
    field_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = field_table.rows[0].cells
    hdr_cells[0].text = "Field Name"
    hdr_cells[1].text = "Extracted Declaration"
    hdr_cells[2].text = "Confidence"
    hdr_cells[3].text = "Rule Status"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    for f in fields_list:
        row_cells = field_table.add_row().cells
        detected = f.get('detected', False)
        status_text = "PASS" if detected else "MISSING"
        if f.get('is_manually_corrected'):
            status_text += " (Edited)"

        row_cells[0].text = f.get('display_name') or f.get('field_id', '').replace('_', ' ').title()
        row_cells[1].text = f.get('value') or "Declaration absent on package"
        row_cells[2].text = f"{int(f.get('confidence', 0) * 100)}%"
        row_cells[3].text = status_text

        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 2. Itemized Violations
    violations_list = scan_data.get('violations', [])
    h2 = doc.add_heading(f"2. Statutory Violations & Non-Compliance Findings ({len(violations_list)})", level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)

    if not violations_list:
        p_nov = doc.add_paragraph("No statutory violations detected. Package declarations comply with Legal Metrology Act 2009.")
        p_nov.runs[0].font.size = Pt(9.5)
        p_nov.runs[0].italic = True
    else:
        v_table = doc.add_table(rows=1, cols=3)
        v_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        v_hdr = v_table.rows[0].cells
        v_hdr[0].text = "Rule Reference"
        v_hdr[1].text = "Severity"
        v_hdr[2].text = "Violation Details & Statutory Remedy"
        for cell in v_hdr:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for v in violations_list:
            v_cells = v_table.add_row().cells
            v_cells[0].text = v.get('rule_code', 'LMPC 2011')
            v_cells[1].text = (v.get('severity') or 'MINOR').upper()

            desc = f"{v.get('title', '')}\n{v.get('description', '')}"
            if v.get('recommendation'):
                desc += f"\nMandatory Action: {v.get('recommendation')}"
            v_cells[2].text = desc

            for cell in v_cells:
                cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Official Sign-off
    p_sign = doc.add_paragraph()
    p_sign.add_run("_____________________________________\n").bold = True
    p_sign.add_run("Authorised Legal Metrology Officer\n").bold = True
    p_sign.add_run("State Enforcement Division, Legal Metrology")

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()

    if output_path:
        with open(output_path, 'wb') as f:
            f.write(docx_bytes)

    return docx_bytes


def generate_excel_report(scan_data: dict, output_path: Optional[str] = None) -> bytes:
    """Generate an Excel workbook summary for this scan.

    Args:
        scan_data: Dict containing scan details, extracted_fields, and violations.
        output_path: Optional file path to save the generated xlsx.

    Returns:
        bytes of the generated Excel workbook.
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        return _generate_csv_fallback(scan_data).encode('utf-8')

    wb = openpyxl.Workbook()

    # Sheet 1: Summary & Fields
    ws_summary = wb.active
    ws_summary.title = "Compliance Audit"

    header_font = Font(name='Arial', size=14, bold=True, color='1E3A8A')
    section_font = Font(name='Arial', size=11, bold=True, color='0F172A')
    bold_font = Font(name='Arial', size=10, bold=True)
    regular_font = Font(name='Arial', size=10)

    header_fill = PatternFill(start_color='E2E8F0', end_color='E2E8F0', fill_type='solid')
    pass_fill = PatternFill(start_color='DCFCE7', end_color='DCFCE7', fill_type='solid')
    fail_fill = PatternFill(start_color='FEE2E2', end_color='FEE2E2', fill_type='solid')

    ws_summary.append(["LEGAL METROLOGY COMPLIANCE AUDIT REPORT"])
    ws_summary.cell(row=1, column=1).font = header_font
    ws_summary.append([])

    # Metadata
    meta = [
        ["Inspection ID", f"#{scan_data.get('id', 'N/A')}", "Status", (scan_data.get('status') or '').upper()],
        ["Product Name", scan_data.get('product_name') or 'N/A', "Compliance Score", f"{scan_data.get('compliance_score', 0):.1f}%"],
        ["Brand", scan_data.get('brand') or 'Unspecified', "Category", (scan_data.get('category') or '').capitalize()],
        ["Audit Date", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"), "Inspector", scan_data.get('user', {}).get('full_name') or 'Inspector'],
    ]
    for row in meta:
        ws_summary.append(row)

    ws_summary.append([])
    ws_summary.append(["MANDATORY DECLARATIONS CHECKLIST (Rule 6)"])
    ws_summary.cell(row=ws_summary.max_row, column=1).font = section_font

    ws_summary.append(["Field Name", "Extracted Value", "Confidence", "Status", "Manual Correction"])
    for col in range(1, 6):
        cell = ws_summary.cell(row=ws_summary.max_row, column=col)
        cell.font = bold_font
        cell.fill = header_fill

    for f in scan_data.get('extracted_fields', []):
        detected = f.get('detected', False)
        status_text = "PASS" if detected else "MISSING"
        ws_summary.append([
            f.get('display_name') or f.get('field_id'),
            f.get('value') or '—',
            f"{int(f.get('confidence', 0) * 100)}%",
            status_text,
            "Yes" if f.get('is_manually_corrected') else "No"
        ])
        status_cell = ws_summary.cell(row=ws_summary.max_row, column=4)
        status_cell.fill = pass_fill if detected else fail_fill

    # Sheet 2: Violations
    ws_violations = wb.create_sheet(title="Violations & Actions")
    ws_violations.append(["Rule Code", "Severity", "Violation Title", "Description", "Statutory Action"])
    for col in range(1, 6):
        cell = ws_violations.cell(row=1, column=col)
        cell.font = bold_font
        cell.fill = header_fill

    for v in scan_data.get('violations', []):
        ws_violations.append([
            v.get('rule_code', 'LMPC'),
            v.get('severity', 'MINOR'),
            v.get('title', ''),
            v.get('description', ''),
            v.get('recommendation', '')
        ])

    # Auto-adjust column widths
    for sheet in [ws_summary, ws_violations]:
        for col in sheet.columns:
            max_len = max(len(str(cell.value or '')) for cell in col)
            col_letter = openpyxl.utils.get_column_letter(col[0].column)
            sheet.column_dimensions[col_letter].width = max(max_len + 3, 12)

    buffer = io.BytesIO()
    wb.save(buffer)
    excel_bytes = buffer.getvalue()
    buffer.close()

    if output_path:
        with open(output_path, 'wb') as f:
            f.write(excel_bytes)

    return excel_bytes


def generate_monthly_ledger_excel(scans: list[dict], month_year: str, output_path: Optional[str] = None) -> bytes:
    """Generate official Month-End Regulatory Activity Ledger Excel report.

    Args:
        scans: List of scan summary/detail dictionaries for the month.
        month_year: Month string e.g. "2026-08".
        output_path: Optional destination path.

    Returns:
        bytes of generated Excel ledger.
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        csv_out = f"Monthly Activity Ledger - {month_year}\nScan ID,Product,Category,Score,Verdict,Location,Sanction\n"
        for s in scans:
            csv_out += f"{s.get('id')},{s.get('product_name')},{s.get('category')},{s.get('compliance_score')},{s.get('status')},{s.get('location_name')},{s.get('approval_status')}\n"
        return csv_out.encode('utf-8')

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Ledger {month_year}"

    header_font = Font(name='Arial', size=13, bold=True, color='1E3A8A')
    bold_font = Font(name='Arial', size=10, bold=True)
    header_fill = PatternFill(start_color='E2E8F0', end_color='E2E8F0', fill_type='solid')
    pass_fill = PatternFill(start_color='DCFCE7', end_color='DCFCE7', fill_type='solid')
    fail_fill = PatternFill(start_color='FEE2E2', end_color='FEE2E2', fill_type='solid')

    ws.append([f"DIRECTORATE OF LEGAL METROLOGY — MONTHLY INSPECTORATE LEDGER ({month_year})"])
    ws.cell(row=1, column=1).font = header_font
    ws.append([])

    # Table Header
    headers = [
        "Scan ID", "Inspection Timestamp", "Auditing Officer", "Product Name",
        "Category", "GPS Location", "SHA-256 Digest", "Score", "Verdict", "Executive Sanction"
    ]
    ws.append(headers)
    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=3, column=col_idx)
        cell.font = bold_font
        cell.fill = header_fill

    for s in scans:
        status_val = (s.get('status') or '').upper()
        ws.append([
            f"#{s.get('id')}",
            s.get('created_at', datetime.now(timezone.utc).strftime("%Y-%m-%d")),
            (s.get('user') or {}).get('full_name') or 'Field Inspector',
            s.get('product_name') or 'Standard Packaging',
            (s.get('category') or 'General').capitalize(),
            s.get('location_name') or 'Field Tagged',
            (s.get('client_evidence_hash') or 'E3B0C442...')[:16] + '...',
            f"{s.get('compliance_score', 0):.1f}%",
            status_val,
            (s.get('approval_status') or 'auto_approved').upper(),
        ])
        status_cell = ws.cell(row=ws.max_row, column=9)
        if status_val == "COMPLIANT":
            status_cell.fill = pass_fill
        elif status_val == "NON_COMPLIANT":
            status_cell.fill = fail_fill

    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = openpyxl.utils.get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 3, 14)

    buffer = io.BytesIO()
    wb.save(buffer)
    excel_bytes = buffer.getvalue()
    buffer.close()

    if output_path:
        with open(output_path, 'wb') as f:
            f.write(excel_bytes)

    return excel_bytes


def _generate_csv_fallback(scan_data: dict) -> str:
    """Generate simple CSV fallback if openpyxl is not available."""
    lines = [
        f"Inspection ID,{scan_data.get('id', '')}",
        f"Product,{scan_data.get('product_name', '')}",
        f"Status,{scan_data.get('status', '')}",
        f"Score,{scan_data.get('compliance_score', 0)}",
        "",
        "Field Name,Value,Detected,Confidence",
    ]
    for f in scan_data.get('extracted_fields', []):
        lines.append(f"{f.get('display_name', '')},{f.get('value', '')},{f.get('detected', '')},{f.get('confidence', '')}")
    return "\n".join(lines)


def generate_pre_market_clearance_docx(pm_data: dict) -> bytes:
    """Generate official Directorate Pre-Market Clearance Certificate in Word (.docx) format."""
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return f"PRE-MARKET CLEARANCE CERTIFICATE: {pm_data.get('certificate_number')}".encode('utf-8')

    doc = docx.Document()

    # Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_gov = p_title.add_run("GOVERNMENT OF INDIA • MINISTRY OF CONSUMER AFFAIRS\nDIRECTORATE OF LEGAL METROLOGY\n")
    run_gov.bold = True
    run_gov.font.size = Pt(14)
    run_gov.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    run_cert = p_title.add_run("CERTIFICATE OF PRE-MARKET PACKAGING CLEARANCE\n")
    run_cert.bold = True
    run_cert.font.size = Pt(13)
    run_cert.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)

    cert_no = pm_data.get('certificate_number') or f"LMPC/PMC/{datetime.now().year}/0091"
    run_no = p_title.add_run(f"Registration Certificate No: {cert_no}\n")
    run_no.font.size = Pt(10)
    run_no.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Details Table
    meta_table = doc.add_table(rows=4, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = True

    meta_data = [
        ("Applicant Enterprise:", pm_data.get('company_name') or 'Brand Owner', "GSTIN / FSSAI ID:", pm_data.get('gstin_fssai_id') or 'Registered'),
        ("Commodity Name:", pm_data.get('product_name') or 'Packaging Commodity', "Brand / Trademark:", pm_data.get('brand') or 'Brand'),
        ("Declared Net Quantity:", str(pm_data.get('declared_net_quantity') or 'N/A'), "Declared MRP:", f"₹{pm_data.get('declared_mrp', '0.00')} (incl. of all taxes)"),
        ("Industry Category:", (pm_data.get('category') or 'General').capitalize(), "Packaging Type:", pm_data.get('packaging_type') or 'Pouch/Box'),
    ]

    for row_idx, row_values in enumerate(meta_data):
        for col_idx, val in enumerate(row_values):
            cell = meta_table.cell(row_idx, col_idx)
            cell.text = val
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx % 2 == 0:
                p.runs[0].bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Statutory Findings
    h = doc.add_heading("Statutory Rule 6 & Schedule II Verification Findings", level=2)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(6)

    findings = [
        ("Rule 6(1)(a) Generic Name Declaration", "Verified & Compliant"),
        ("Rule 6(1)(b) Standardized SI Net Quantity", "Verified & Compliant"),
        ("Rule 6(1)(c) Maximum Retail Price (MRP)", "Verified & Compliant"),
        ("Rule 9 / Schedule II Font Height Proportions", "Verified & Compliant"),
        ("Consumer Care Details & Packer Address", "Verified & Compliant"),
    ]
    f_table = doc.add_table(rows=1, cols=2)
    f_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    f_table.rows[0].cells[0].text = "Statutory Declaration Requirement"
    f_table.rows[0].cells[1].text = "Verification Status"
    f_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    f_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    for req, res in findings:
        r = f_table.add_row().cells
        r[0].text = req
        r[1].text = res

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Remarks
    p_rem = doc.add_paragraph()
    p_rem.add_run(f"Directorate Approval Remarks: \"{pm_data.get('supervisor_notes') or 'Pre-market packaging design clearance granted for commercial printing and distribution.'}\"")
    p_rem.runs[0].italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes


def generate_pre_market_clearance_excel(pm_data: dict) -> bytes:
    """Generate official Directorate Pre-Market Clearance Record in Excel (.xlsx) format."""
    try:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    except ImportError:
        return f"Certificate No: {pm_data.get('certificate_number')}".encode('utf-8')

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Packaging Clearance Certificate"

    # Header
    ws.merge_cells("A1:D1")
    ws["A1"] = "DIRECTORATE OF LEGAL METROLOGY — PRE-MARKET PACKAGING CLEARANCE CERTIFICATE"
    ws["A1"].font = Font(name="Calibri", size=14, bold=True, color="FFFFFF")
    ws["A1"].fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36

    cert_no = pm_data.get('certificate_number') or f"LMPC/PMC/{datetime.now().year}/0091"
    ws.merge_cells("A2:D2")
    ws["A2"] = f"Official Certificate No: {cert_no} | Status: APPROVED & CERTIFIED"
    ws["A2"].font = Font(name="Calibri", size=11, bold=True, color="065F46")
    ws["A2"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 24

    fields = [
        ("Applicant Enterprise", pm_data.get('company_name') or 'Brand Owner'),
        ("GSTIN / FSSAI License", pm_data.get('gstin_fssai_id') or 'Registered'),
        ("Commodity Name", pm_data.get('product_name') or 'Packaging Commodity'),
        ("Brand / Trademark", pm_data.get('brand') or 'N/A'),
        ("Declared MRP (₹)", f"₹{pm_data.get('declared_mrp', '0.00')} (incl. of all taxes)"),
        ("Declared Net Quantity", str(pm_data.get('declared_net_quantity') or 'N/A')),
        ("Industry Sector", (pm_data.get('category') or 'General').capitalize()),
        ("Packaging Format", pm_data.get('packaging_type') or 'Packet / Pouch'),
        ("Statutory Rule 6 Status", "100% PASS — All Mandatory Declarations Present"),
        ("Schedule II Font Size", "COMPLIANT — Minimum Height Threshold Satisfied"),
        ("Supervisor Remarks", pm_data.get('supervisor_notes') or 'Pre-market packaging design clearance granted.'),
        ("Clearance Granted Timestamp", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
    ]

    for idx, (label, val) in enumerate(fields, start=4):
        ws[f"A{idx}"] = label
        ws[f"A{idx}"].font = Font(name="Calibri", size=10, bold=True)
        ws.merge_cells(f"B{idx}:D{idx}")
        ws[f"B{idx}"] = val
        ws[f"B{idx}"].font = Font(name="Calibri", size=10)

    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 25
    ws.column_dimensions["C"].width = 25
    ws.column_dimensions["D"].width = 25

    buffer = io.BytesIO()
    wb.save(buffer)
    excel_bytes = buffer.getvalue()
    buffer.close()
    return excel_bytes


